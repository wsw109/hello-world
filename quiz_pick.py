#!/usr/bin/python3

import docx
import re
#from quiz_selection import quiz_selection
import random

# pick 25 randomly from different types of questions.
# No.1 ~ No.10   Architecture rules
# No.11 ~ No.40  Coding rules
# No.51 ~ No.90  C++11 new features
# No.91 ~ No.100 C++14 new features

#examples:
#arc_rule_num = 5
#code_rule_num = 7
#cpp11_num = 10
#cpp14_num = 3

def quiz_selection(arc_rule_num, code_rule_num, cpp11_num, cpp14_num):
	selected = random.sample(range(1,10),arc_rule_num) \
	    + random.sample(range(11,40),code_rule_num) \
	    + random.sample(range(41,90),cpp11_num) \
	    + random.sample(range(91,100),cpp14_num)
	selected.sort()
	return selected

#read document
file=docx.Document("code_quiz.docx")
#print("paragraphs:"+str(len(file.paragraphs))) 
file_word = docx.Document()

#for para in file.paragraphs:
#    print(para.text)

dict
dict_num = 0
dict = {dict_num: ''}

for i in range(len(file.paragraphs)):
    t = file.paragraphs[i].text
    s = re.match(r'^\d+\.\s+',t)
    if s:
        dict_num += 1
        dict[dict_num] = t
    else:
        dict[dict_num] += '\n' + t
    #print(dict_num)
    #print(dict[dict_num])

print('dict_num: ',dict_num)

#print quiz title and attention
file_word.add_paragraph(dict[0])

#randomly generated quiz questions
print('Randomly generated quiz questions:')
selected = quiz_selection(5,7,10,3)
print(selected)

#question_num is the new generated questions number
j=1
for key in selected:
    #print(dict[key])
    #print('\n')
    question_num = str(j) + '. '
    file_word.add_paragraph(question_num + dict[key][3:])
    j += 1

#create the selected quiz file
file_word.save("quiz_selected.docx")

